# -*- coding: utf-8 -*-
"""
DOCX Document Generator
python-docx를 사용한 2025 개정 표준 준수 문서 생성
"""
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


class DOCXGenerator:
    """DOCX 문서 생성기 (2025 개정 표준 준수)"""
    
    def __init__(self):
        """초기화"""
        self.font_name = "바탕"  # 기본 폰트
        self.title_size = Pt(16)
        self.body_size = Pt(11)
        self.line_spacing = 1.15
        
    def generate_official_document(self, doc_data: dict) -> BytesIO:
        """
        공문서 DOCX 생성
        
        Args:
            doc_data: 문서 데이터
                - title: 문서 제목
                - receiver: 수신자
                - body_paragraphs: 본문 단락
                - doc_num: 문서번호
                - doc_date: 시행일자
                - department_head: 발신명의
        
        Returns:
            BytesIO: DOCX 파일 바이트 스트림
        """
        # 데이터 검증
        if not isinstance(doc_data, dict):
            raise ValueError(f"doc_data must be dict, got {type(doc_data)}")
        
        doc = Document()
        
        # 페이지 설정
        self._set_page_margins(doc)
        
        # 직인생략 (우측 상단)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("직인생략")
        self._set_font(run, size=Pt(9), color=RGBColor(128, 128, 128))
        
        # 제목
        title = doc_data.get("title", "공문서")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        self._set_font(run, size=self.title_size, bold=True)
        
        # 문서정보 (문서번호, 시행일자, 수신)
        doc.add_paragraph()  # 빈 줄
        
        # 문서번호
        doc_num = doc_data.get("doc_num", self._generate_doc_num())
        p = doc.add_paragraph()
        run = p.add_run(f"문서번호: {doc_num}")
        self._set_font(run, size=self.body_size)
        
        # 시행일자
        doc_date = self._format_date_standard(doc_data.get("doc_date"))
        p = doc.add_paragraph()
        run = p.add_run(f"시행일자: {doc_date}")
        self._set_font(run, size=self.body_size)
        
        # 수신
        receiver = doc_data.get("receiver", "수신자")
        p = doc.add_paragraph()
        run = p.add_run(f"수    신: {receiver}")
        self._set_font(run, size=self.body_size)
        
        # 구분선
        doc.add_paragraph("─" * 50)
        
        # 본문
        doc.add_paragraph()  # 빈 줄
        body_paragraphs = doc_data.get("body_paragraphs", [])
        if isinstance(body_paragraphs, str):
            body_paragraphs = body_paragraphs.split("\n")
        
        for para_text in body_paragraphs:
            if not para_text or not para_text.strip():
                continue
            
            # 항목 체계 적용
            indent_level, formatted_text = self._detect_hierarchy(para_text.strip())
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(indent_level * 0.25)  # 2타 = 0.25인치
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = self.line_spacing
            
            run = p.add_run(formatted_text)
            self._set_font(run, size=self.body_size)
        
        # 끝. 추가
        if body_paragraphs:
            last_p = doc.paragraphs[-1]
            if not last_p.text.endswith("끝."):
                last_p.add_run(" 끝.")
        
        # 발신명의
        doc.add_paragraph()  # 빈 줄
        doc.add_paragraph()  # 빈 줄
        
        department_head = doc_data.get("department_head", "행정기관장")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(department_head)
        self._set_font(run, size=self.body_size, bold=True)
        
        # BytesIO로 저장
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    
    def generate_processing_guide(self, workflow_data: dict) -> BytesIO:
        """
        처리가이드 보고서 DOCX 생성
        
        Args:
            workflow_data: 워크플로우 데이터
                - analysis: 분석 결과
                - law: 법적 검토
                - strategy: 처리 전략
                - procedure: 절차
        
        Returns:
            BytesIO: DOCX 파일 바이트 스트림
        """
        # 데이터 검증 및 변환
        if isinstance(workflow_data, str):
            try:
                import json
                workflow_data = json.loads(workflow_data)
            except:
                # JSON 파싱 실패 시 기본 구조 생성
                workflow_data = {
                    "analysis": {
                        "summary": [str(workflow_data)],
                        "case_type": "일반 민원"
                    }
                }
        
        if not isinstance(workflow_data, dict):
            # 최후의 수단
            workflow_data = {
                "analysis": {
                    "summary": [str(workflow_data)]
                }
            }
        
        doc = Document()
        
        # 페이지 설정
        self._set_page_margins(doc)
        
        # 제목
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("행정 처리 가이드")
        self._set_font(run, size=Pt(18), bold=True)
        
        # 작성일
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"작성일: {self._format_date_standard(None)}")
        self._set_font(run, size=Pt(9), color=RGBColor(100, 100, 100))
        
        doc.add_paragraph()  # 빈 줄
        
        # 데이터 추출
        # 데이터 추출 및 타입 안전성 확보
        analysis = workflow_data.get("analysis")
        if isinstance(analysis, str):
            analysis = {"summary": [analysis], "case_type": "일반 민원"}
        elif not isinstance(analysis, dict):
            analysis = {}
            
        law = workflow_data.get("law")
        if isinstance(law, str):
            law = {"statutes": [law]}
        elif not isinstance(law, dict):
            law = {}
            
        strategy = workflow_data.get("strategy")
        if isinstance(strategy, str):
            strategy = {"summary": strategy}
        elif not isinstance(strategy, dict):
            strategy = {}
            
        procedure = workflow_data.get("procedure")
        if isinstance(procedure, str):
            procedure = {"checklist": [procedure]}
        elif not isinstance(procedure, dict):
            procedure = {}
        
        # 1. 민원 개요
        self._add_section_header(doc, "1. 민원 개요")
        
        # 케이스 유형
        case_type = analysis.get("case_type", "일반 민원")
        p = doc.add_paragraph()
        run = p.add_run(f"케이스 유형: {case_type}")
        self._set_font(run, size=self.body_size, bold=True)
        
        # 핵심 쟁점
        summary = analysis.get("summary", [])
        if summary:
            self._add_subsection(doc, "핵심 쟁점:")
            if isinstance(summary, str):
                summary = [summary]
            for item in summary:
                self._add_bullet_item(doc, item)
        
        # 리스크 요소
        risks = analysis.get("risks", [])
        if risks:
            self._add_subsection(doc, "리스크 요소:")
            for risk in risks:
                self._add_bullet_item(doc, risk, bullet="⚠")
        
        doc.add_paragraph()  # 빈 줄
        
        # 2. 법적 근거
        self._add_section_header(doc, "2. 법적 근거")
        
        # 관련 법령
        statutes = law.get("statutes", [])
        if statutes:
            self._add_subsection(doc, "관련 법령:")
            for statute in statutes:
                self._add_bullet_item(doc, statute)
        
        # 판례
        precedents = law.get("precedents", [])
        if precedents:
            self._add_subsection(doc, "판례:")
            for precedent in precedents:
                self._add_bullet_item(doc, precedent)
        
        if not statutes and not precedents:
            p = doc.add_paragraph()
            run = p.add_run("법적 검토 내용 없음")
            self._set_font(run, size=self.body_size, italic=True)
        
        doc.add_paragraph()  # 빈 줄
        
        # 3. 검토 의견
        self._add_section_header(doc, "3. 검토 의견")
        
        # 처리 방안 요약
        strategy_summary = strategy.get("summary", "")
        if strategy_summary:
            p = doc.add_paragraph()
            run = p.add_run(f"처리 방안: {strategy_summary}")
            self._set_font(run, size=self.body_size)
        
        # 단계별 조치 계획
        steps = strategy.get("steps", [])
        if steps:
            self._add_subsection(doc, "단계별 조치 계획:")
            if isinstance(steps, str):
                steps = [steps]
            for i, step in enumerate(steps, 1):
                self._add_numbered_item(doc, i, step)
        
        doc.add_paragraph()  # 빈 줄
        
        # 4. 향후 계획
        self._add_section_header(doc, "4. 향후 계획")
        
        # 체크리스트
        checklist = procedure.get("checklist", [])
        if checklist:
            self._add_subsection(doc, "체크리스트:")
            for item in checklist:
                self._add_bullet_item(doc, item, bullet="☐")
        
        # 필요 서식
        templates = procedure.get("templates", [])
        if templates:
            self._add_subsection(doc, "필요 서식:")
            for template in templates:
                self._add_bullet_item(doc, template)
        
        # 다음 조치
        next_action = procedure.get("next_action", "")
        if next_action:
            self._add_subsection(doc, "다음 조치:")
            p = doc.add_paragraph()
            run = p.add_run(next_action)
            self._set_font(run, size=self.body_size)
        
        # BytesIO로 저장
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    
    def _set_page_margins(self, doc):
        """페이지 여백 설정"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
    
    def _set_font(self, run, size=None, bold=False, italic=False, color=None):
        """폰트 설정"""
        run.font.name = self.font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
        
        if size:
            run.font.size = size
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
        if color:
            run.font.color.rgb = color
    
    def _detect_hierarchy(self, text: str):
        """
        항목 체계 감지 및 들여쓰기 레벨 반환
        
        Returns:
            (indent_level, formatted_text)
        """
        patterns = [
            (r'^(\d+)\.\s', 0),      # 1. (대항목, 0타)
            (r'^([가-힣])\.\s', 1),  # 가. (중항목, 2타 = 1레벨)
            (r'^(\d+)\)\s', 2),      # 1) (소항목, 4타 = 2레벨)
            (r'^([가-힣])\)\s', 3),  # 가) (세부항목, 6타 = 3레벨)
            (r'^\((\d+)\)\s', 4),    # (1) (5레벨)
            (r'^\(([가-힣])\)\s', 5),# (가) (6레벨)
        ]
        
        for pattern, level in patterns:
            if re.match(pattern, text):
                return level, text
        
        return 0, text
    
    def _format_date_standard(self, date_input) -> str:
        """2025 개정 표준 날짜 형식"""
        if date_input is None:
            dt = datetime.now()
        elif isinstance(date_input, str):
            if ". " in date_input:
                return date_input
            try:
                dt = datetime.fromisoformat(date_input.replace("Z", "+00:00"))
            except:
                dt = datetime.now()
        elif isinstance(date_input, datetime):
            dt = date_input
        else:
            dt = datetime.now()
        
        # 2025 개정 표준: "2025. 1. 27."
        return f"{dt.year}. {dt.month}. {dt.day}."
    
    def _generate_doc_num(self) -> str:
        """문서번호 자동 생성"""
        today = datetime.now()
        return f"행정-{today.year}-{today.strftime('%m%d')}-001호"
    
    def _add_section_header(self, doc, text):
        """섹션 헤더 추가"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        self._set_font(run, size=Pt(13), bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    
    def _add_subsection(self, doc, text):
        """서브섹션 추가"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        self._set_font(run, size=self.body_size, bold=True)
        p.paragraph_format.space_before = Pt(6)
    
    def _add_bullet_item(self, doc, text, bullet="•"):
        """불렛 아이템 추가"""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = self.line_spacing
        
        run = p.add_run(f"{bullet} {text}")
        self._set_font(run, size=self.body_size)
    
    def _add_numbered_item(self, doc, number, text):
        """번호 아이템 추가"""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = self.line_spacing
        
        run = p.add_run(f"{number}. {text}")
        self._set_font(run, size=self.body_size)


# 편의 함수
def generate_official_docx(doc_data: dict) -> BytesIO:
    """공문서 DOCX 생성 (편의 함수)"""
    generator = DOCXGenerator()
    return generator.generate_official_document(doc_data)


def generate_guide_docx(workflow_data: dict) -> BytesIO:
    """처리가이드 DOCX 생성 (편의 함수)"""
    generator = DOCXGenerator()
    return generator.generate_processing_guide(workflow_data)
