# -*- coding: utf-8 -*-
"""
Govable AI - DOCX ë¬¸ì„œ ìƒì„± ëª¨ë“ˆ
python-docxë¥¼ ì‚¬ìš©í•œ ê³µë¬¸ì„œ ë° ë³´ê³ ì„œ ìƒì„±
"""
import io
from typing import Dict, Any, List, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None


def generate_official_docx(doc_data: Dict[str, Any]) -> bytes:
    """
    ê³µë¬¸ì„œ DOCX ìƒì„±
    
    Args:
        doc_data: ë¬¸ì„œ ë°ì´í„°
            - title: ë¬¸ì„œ ì œëª©
            - receiver: ìˆ˜ì‹ ì (ì„ íƒ)
            - body_paragraphs: ë³¸ë¬¸ ë¬¸ë‹¨ ë¦¬ìŠ¤íŠ¸
            - department_head: ë¶€ì„œì¥ ì´ë¦„ (ì„ íƒ)
            - doc_num: ë¬¸ì„œë²ˆí˜¸ (ì„ íƒ)
            
    Returns:
        DOCX íŒŒì¼ ë°”ì´íŠ¸
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx íŒ¨í‚¤ì§€ê°€ ì„¤ì¹˜ë˜ì§€ ì•Šì•˜ìŠµë‹ˆë‹¤. pip install python-docx")
    
    doc = Document()
    
    # ë¬¸ì„œ ì œëª©
    title = doc_data.get('title', 'ê³µë¬¸ì„œ')
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ë¬¸ì„œ ì •ë³´
    doc_num = doc_data.get('doc_num', '')
    today = datetime.now().strftime('%Yë…„ %mì›” %dì¼')
    
    if doc_num:
        p = doc.add_paragraph()
        p.add_run(f"ë¬¸ì„œë²ˆí˜¸: {doc_num}").bold = True
    
    p = doc.add_paragraph()
    p.add_run(f"ì‹œí–‰ì¼ì: {today}").bold = True
    
    receiver = doc_data.get('receiver', '')
    if receiver:
        p = doc.add_paragraph()
        p.add_run(f"ìˆ˜ì‹ : {receiver}").bold = True
    
    doc.add_paragraph()  # ë¹ˆ ì¤„
    
    # ë³¸ë¬¸
    body_paragraphs = doc_data.get('body_paragraphs', [])
    if isinstance(body_paragraphs, str):
        body_paragraphs = body_paragraphs.split('\n\n')
    
    for para_text in body_paragraphs:
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            p.style = 'Normal'
    
    doc.add_paragraph()  # ë¹ˆ ì¤„
    
    # ë°œì‹  ì •ë³´
    department_head = doc_data.get('department_head', '')
    if department_head:
        p = doc.add_paragraph()
        p.add_run(f"{department_head}").bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # ë°”ì´íŠ¸ë¡œ ë³€í™˜
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_guide_docx(guide_data: Dict[str, Any]) -> bytes:
    """
    ì²˜ë¦¬ê°€ì´ë“œ ë³´ê³ ì„œ DOCX ìƒì„±
    
    Args:
        guide_data: ê°€ì´ë“œ ë°ì´í„°
            - analysis: ë¶„ì„ ê²°ê³¼ ë”•ì…”ë„ˆë¦¬
            - summary: ìš”ì•½ (ì„ íƒ)
            - timeline: ì²˜ë¦¬ ì ˆì°¨ (ì„ íƒ)
            - legal_basis: ë²•ì  ê·¼ê±° (ì„ íƒ)
            
    Returns:
        DOCX íŒŒì¼ ë°”ì´íŠ¸
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx íŒ¨í‚¤ì§€ê°€ ì„¤ì¹˜ë˜ì§€ ì•Šì•˜ìŠµë‹ˆë‹¤. pip install python-docx")
    
    doc = Document()
    
    # ì œëª©
    title = guide_data.get('title', 'ë¯¼ì› ì²˜ë¦¬ ê°€ì´ë“œ')
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ì‘ì„±ì¼
    today = datetime.now().strftime('%Yë…„ %mì›” %dì¼')
    p = doc.add_paragraph()
    p.add_run(f"ì‘ì„±ì¼: {today}").italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()  # ë¹ˆ ì¤„
    
    # ë¶„ì„ ê²°ê³¼
    analysis = guide_data.get('analysis', {})
    
    # ìš”ì•½
    if 'summary' in analysis or 'summary' in guide_data:
        doc.add_heading('ğŸ“‹ ìš”ì•½', level=2)
        summary = analysis.get('summary', guide_data.get('summary', ''))
        if summary:
            doc.add_paragraph(str(summary))
        doc.add_paragraph()
    
    # ì²˜ë¦¬ ì ˆì°¨
    if 'timeline' in analysis or 'timeline' in guide_data:
        doc.add_heading('ğŸ“… ì²˜ë¦¬ ì ˆì°¨', level=2)
        timeline = analysis.get('timeline', guide_data.get('timeline', []))
        
        if isinstance(timeline, list):
            for i, step in enumerate(timeline, 1):
                if isinstance(step, dict):
                    step_text = step.get('description', str(step))
                else:
                    step_text = str(step)
                doc.add_paragraph(f"{i}. {step_text}", style='List Number')
        else:
            doc.add_paragraph(str(timeline))
        doc.add_paragraph()
    
    # ë²•ì  ê·¼ê±°
    if 'legal_basis' in analysis or 'legal_basis' in guide_data:
        doc.add_heading('âš–ï¸ ë²•ì  ê·¼ê±°', level=2)
        legal = analysis.get('legal_basis', guide_data.get('legal_basis', ''))
        if legal:
            doc.add_paragraph(str(legal))
        doc.add_paragraph()
    
    # ì²˜ë¦¬ ì „ëµ
    if 'strategy' in analysis or 'strategy' in guide_data:
        doc.add_heading('ğŸ¯ ì²˜ë¦¬ ì „ëµ', level=2)
        strategy = analysis.get('strategy', guide_data.get('strategy', ''))
        if strategy:
            doc.add_paragraph(str(strategy))
        doc.add_paragraph()
    
    # ê¸°íƒ€ í•„ë“œë“¤ (analysis ë”•ì…”ë„ˆë¦¬ì˜ ë‚˜ë¨¸ì§€)
    known_fields = {'summary', 'timeline', 'legal_basis', 'strategy', 'title'}
    for key, value in analysis.items():
        if key not in known_fields and value:
            # keyë¥¼ ë¬¸ìì—´ë¡œ ë³€í™˜ í›„ ì²˜ë¦¬
            key_str = str(key).replace("_", " ").title()
            doc.add_heading(f'ğŸ“Œ {key_str}', level=2)
            if isinstance(value, list):
                for item in value:
                    doc.add_paragraph(f"â€¢ {item}", style='List Bullet')
            else:
                doc.add_paragraph(str(value))
            doc.add_paragraph()
    
    # ë°”ì´íŠ¸ë¡œ ë³€í™˜
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
